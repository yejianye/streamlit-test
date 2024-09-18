import streamlit as st
import os
import io

from docx import Document
from docx.shared import Pt
from openai import OpenAI
from pprint import pprint
# from joblib import Memory

DEFAULT_FONT = 'Microsoft YaHei'
DEFAULT_MODEL = 'gpt-4o'

# mem = Memory(location='/tmp/openai_cache')
# @mem.cache
def llm_completion(prompt):
    client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
    resp = client.chat.completions.create(
        model=DEFAULT_MODEL,
        messages=[{"role": "user", "content": prompt}],
    )
    return resp.choices[0].message.content

def get_doc_content(doc):
    text = [para.text for para in doc.paragraphs]

    return '\n'.join(text)

def get_highlighted_words(doc):
    highlighted_words = []

    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.highlight_color:
                highlighted_words.append(run.text)

    return highlighted_words

def translate_to_chinese(text):
    prompt = f"""
    Translate INPUT_TEXT to Chinese. Do not output anything other than the translation.

    INPUT_TEXT:
    {text}
    """
    return llm_completion(prompt)

def create_vocabulary(words, article):
    words = '\n'.join(words)
    prompt = f"""
    对于 WORDS 中的每一个单词或是词组，在 ARTICLE 的上下文语境中翻译。每个单词输出一行，格式为
    单词 | 音标 | 词性 | 中文翻译

    - 如果是一个词组，则音标与词性可以为空，但 | 不能省略。
    - 音标左右两边需要用 / 来包裹。
    - 除了上述要求的输出格式以外，不要输出任何其他内容。

    # WORDS
    {words}

    # ARTICLE
    {article}
    """
    result = llm_completion(prompt)
    result = result.strip().split('\n')
    result = [[j.strip() for j in i.split('|')] for i in result]
    return result

def add_content_to_doc(doc, content, font_family=DEFAULT_FONT, font_size=10):
    # Split the long_text into paragraphs by newline
    paragraphs = content.split('\n')
    # Loop through each paragraph in the list
    for para_text in paragraphs:
        if para_text.strip():  # Ensure you're not adding empty paragraphs
            # Add a new paragraph to the end of the document
            para = doc.add_paragraph(style=None)
            run = para.add_run(para_text)
            run.font.size = Pt(font_size)

def add_vocabulary_table(doc, vocabulary):
    table = doc.tables[0]
    for word in vocabulary:
        row = table.add_row()
        row.height = Pt(20)
        for j, cell in enumerate(row.cells):
            cell.text = word[j]

def gen_translation_and_vocabulary(input_file, output_file, show_progress=False):
    doc = Document(input_file)
    content = get_doc_content(doc)
    if show_progress:
        st.write("Translating doc to Chinese...")
    cn_content = translate_to_chinese(content)

    words = get_highlighted_words(doc)
    if show_progress:
        st.write("Creating vocabulary from highlighted words...")
    vocabulary = create_vocabulary(words, content)

    output_doc = Document("template.docx")
    add_vocabulary_table(output_doc, vocabulary)
    add_content_to_doc(output_doc, cn_content)
    output_doc.save(output_file)
    if show_progress:
        st.write("The document is ready for download.")

def test_translation():
    doc = Document("test.docx")
    content = get_doc_content(doc)
    print(translate_to_chinese(content))

def test_vocabulary():
    doc = Document("test.docx")
    content = get_doc_content(doc)
    words = get_highlighted_words(doc)
    vocabulary = create_vocabulary(words, content)
    print('\n'.join(words))
    pprint(vocabulary)

def test_add_content():
    doc = Document("template.docx")
    cn_content = """
奥运圣火将在 7 月 26 日点亮这座光之城，届时全球最伟大的体育盛事将在巴黎启动。马术比赛将在凡尔赛宫宏伟的场地举行；排球将在埃菲尔铁塔旁边飞越网格。组织者希望向来访的体育迷、商业高管和外国政要展示法国最佳的一面。成千上万的志愿者中有一位描述了“一种具有感染力的积极能量”。

然而，并非所有当地人都这样兴奋。出于对恐怖主义的担忧，安全措施非常严格，巴黎市中心的许多地区已被划为禁区。这些地区的餐馆和其他商业场所比平常空旷，且七月份最高档的酒店预订量下降了 20%到 50%。这种反应在奥运主办城市中并不少见；没有人喜欢被打扰。但巴黎奥运还激起了其他一些老生常谈的问题：举办奥运会真的值得吗？
    """
    add_content_to_doc(doc, cn_content, font_size=10)
    doc.save("test_add_content.docx")

def test_add_vocabulary_table():
    doc = Document("template.docx")
    voc = [['spectacle', '/ˈspɛk.tə.kəl/', 'n.', '壮观景象'],
           ['under way', '', '', '进行中'],
           ['Dressage', '/drɛˈsɑʒ/', 'n.', '花样骑术'],
           ['Versaille', '/vɛrˈsaɪ/', 'n.', '凡尔赛'],
           ['whizz', '/wɪz/', 'v.', '嗖嗖作响'],
           ['infectious', '/ɪnˈfɛk.ʃəs/', 'adj.', '有感染力的'],
           ['enthused', '/ɪnˈθjuːzd/', 'adj.', '激情洋溢的'],
           ['zoned off', '', '', '划为禁区'],
           ['poshest', '/pɒʃɪst/', 'adj.', '最豪华的']]
    add_vocabulary_table(doc, voc)
    doc.save("test_add_vocabulary.docx")

def main():
    st.title("Create Translation and Vocabulary")
    # File uploader widget
    uploaded_file = st.file_uploader("Choose a Word file")

    if uploaded_file is not None:
        # To read content of the uploaded file
        output_doc = io.BytesIO()
        gen_translation_and_vocabulary(uploaded_file, output_doc, show_progress=True)

        # Display the processed content (optional)

        # Providing a link to download the new file
        st.download_button(label="Download Processed File",
                           data=output_doc.getvalue(),
                           file_name="result.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == "__main__":
    main()
